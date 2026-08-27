from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

import pymupdf

SECTION_ALIASES = {
    "summary": {"summary", "profile", "objective", "about"},
    "experience": {"experience", "work experience", "employment", "professional experience"},
    "education": {"education", "academic", "qualifications"},
    "skills": {"skills", "technical skills", "core competencies", "technologies"},
    "projects": {"projects", "personal projects", "selected projects"},
    "certifications": {"certifications", "certificates", "licenses"},
}

ACTION_VERBS = {
    "led", "built", "created", "designed", "developed", "implemented", "improved",
    "increased", "reduced", "managed", "launched", "optimized", "delivered",
    "automated", "collaborated", "analyzed", "architected", "owned", "shipped",
}


@dataclass
class ParsedResume:
    filename: str
    raw_text: str
    email: str | None = None
    phone: str | None = None
    linkedin: str | None = None
    github: str | None = None
    sections: dict[str, str] = field(default_factory=dict)
    word_count: int = 0
    page_count: int = 1
    bullet_count: int = 0
    quantified_bullets: int = 0
    action_verb_bullets: int = 0


def extract_text(filename: str, data: bytes) -> tuple[str, int]:
    name = filename.lower()
    if name.endswith(".pdf"):
        doc = pymupdf.open(stream=data, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        return text, doc.page_count
    if name.endswith(".docx"):
        from docx import Document

        document = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in document.paragraphs)
        return text, max(1, len(document.paragraphs) // 40)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore"), 1
    raise ValueError("Unsupported file type. Upload PDF, DOCX, or TXT.")


def _first(pattern: str, text: str) -> str | None:
    match = re.search(pattern, text, flags=re.IGNORECASE)
    return match.group(0).strip() if match else None


def _split_sections(text: str) -> dict[str, str]:
    lines = [ln.strip() for ln in text.splitlines()]
    buckets: dict[str, list[str]] = {key: [] for key in SECTION_ALIASES}
    current = "summary"
    for line in lines:
        key = _heading_key(line)
        if key:
            current = key
            continue
        if current in buckets:
            buckets[current].append(line)
    return {k: "\n".join(v).strip() for k, v in buckets.items() if v}


def _heading_key(line: str) -> str | None:
    cleaned = re.sub(r"[^a-zA-Z ]", "", line).strip().lower()
    if not cleaned or len(cleaned.split()) > 4:
        return None
    for key, aliases in SECTION_ALIASES.items():
        if cleaned in aliases:
            return key
    return None


def parse_resume(filename: str, data: bytes) -> ParsedResume:
    raw_text, page_count = extract_text(filename, data)
    raw_text = re.sub(r"\n{3,}", "\n\n", raw_text).strip()
    if not raw_text:
        raise ValueError("No extractable text. Use a text-based file, not a scanned image.")

    bullets = [ln.strip() for ln in raw_text.splitlines() if re.match(r"^[\-\u2022\*]\s+", ln.strip())]
    if not bullets:
        bullets = [ln.strip() for ln in raw_text.splitlines() if ln.strip().startswith("-")]

    quantified = sum(1 for b in bullets if re.search(r"\d", b))
    action = sum(1 for b in bullets if any(b.lower().lstrip("-•* ").startswith(v) for v in ACTION_VERBS))

    return ParsedResume(
        filename=filename,
        raw_text=raw_text,
        email=_first(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", raw_text),
        phone=_first(r"(?:\+?\d{1,3}[\s-]?)?(?:\(?\d{3}\)?[\s-]?)?\d{3}[\s-]?\d{4}", raw_text),
        linkedin=_first(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9_-]+/?", raw_text),
        github=_first(r"(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9_-]+/?", raw_text),
        sections=_split_sections(raw_text),
        word_count=len(re.findall(r"\b\w+\b", raw_text)),
        page_count=page_count,
        bullet_count=len(bullets),
        quantified_bullets=quantified,
        action_verb_bullets=action,
    )
